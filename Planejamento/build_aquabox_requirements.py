from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE

OUT = 'Planejamento_de_Requisitos_Firmware_Aquabox.docx'
BLUE = '2E74B5'; DARK = '1F4D78'; NAVY = '0B2545'; LIGHT = 'E8EEF5'; GRAY = 'F2F4F7'; CALLOUT = 'F4F6F9'

def set_font(run, size=11, bold=None, color=None, italic=None):
    run.font.name = 'Calibri'
    run._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    run.font.size = Pt(size)
    if bold is not None: run.bold = bold
    if italic is not None: run.italic = italic
    if color: run.font.color.rgb = RGBColor.from_string(color)

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'), fill)

def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr(); tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar'); tcPr.append(tcMar)
    for m, v in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node = tcMar.find(qn('w:'+m))
        if node is None: node = OxmlElement('w:'+m); tcMar.append(node)
        node.set(qn('w:w'), str(v)); node.set(qn('w:type'), 'dxa')

def set_table_geometry(table, widths):
    table.autofit = False
    tblPr = table._tbl.tblPr
    tblW = tblPr.first_child_found_in('w:tblW')
    tblW.set(qn('w:w'), '9360'); tblW.set(qn('w:type'), 'dxa')
    ind = OxmlElement('w:tblInd'); ind.set(qn('w:w'), '120'); ind.set(qn('w:type'), 'dxa'); tblPr.append(ind)
    grid = table._tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths): col.set(qn('w:w'), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tcPr=cell._tc.get_or_add_tcPr(); tcW=tcPr.first_child_found_in('w:tcW')
            tcW.set(qn('w:w'), str(width)); tcW.set(qn('w:type'), 'dxa')
            cell.width = Inches(width/1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)

def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    header = OxmlElement('w:tblHeader')
    header.set(qn('w:val'), 'true')
    trPr.append(header)

def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'; table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    hdr = table.rows[0].cells
    for c, text in zip(hdr, headers):
        shade(c, LIGHT); p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(0)
        r=p.add_run(text); set_font(r, 9.5, True, NAVY)
    for i, row in enumerate(rows):
        cells=table.add_row().cells
        for c, text in zip(cells,row):
            p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.05
            r=p.add_run(text); set_font(r, 9.3)
            if i % 2: shade(c, 'FAFBFC')
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table

def add_heading(doc, text, level=1):
    p=doc.add_paragraph(style=f'Heading {level}')
    p.paragraph_format.keep_with_next=True
    p.add_run(text)
    return p

def add_bullets(doc, items):
    for item in items:
        p=doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after=Pt(4); p.paragraph_format.line_spacing=1.15
        p.add_run(item)

def add_callout(doc, title, text):
    table=doc.add_table(rows=1, cols=1); table.alignment=WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table,[9360]); cell=table.cell(0,0); shade(cell,CALLOUT)
    p=cell.paragraphs[0]; p.paragraph_format.space_after=Pt(2)
    r=p.add_run(title); set_font(r,10,True,DARK)
    p=cell.add_paragraph(); p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.1
    r=p.add_run(text); set_font(r,10)
    doc.add_paragraph().paragraph_format.space_after=Pt(0)

doc=Document()
sec=doc.sections[0]
sec.top_margin=Inches(1); sec.bottom_margin=Inches(1); sec.left_margin=Inches(1); sec.right_margin=Inches(1)
sec.header_distance=Inches(.492); sec.footer_distance=Inches(.492)

styles=doc.styles
normal=styles['Normal']; normal.font.name='Calibri'; normal._element.rPr.rFonts.set(qn('w:ascii'),'Calibri'); normal._element.rPr.rFonts.set(qn('w:hAnsi'),'Calibri'); normal.font.size=Pt(11); normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.1
for name,size,color,before,after in [('Heading 1',16,BLUE,16,8),('Heading 2',13,BLUE,12,6),('Heading 3',12,DARK,8,4)]:
    s=styles[name]; s.font.name='Calibri'; s._element.rPr.rFonts.set(qn('w:ascii'),'Calibri'); s._element.rPr.rFonts.set(qn('w:hAnsi'),'Calibri'); s.font.size=Pt(size); s.font.color.rgb=RGBColor.from_string(color); s.font.bold=True; s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after)

# Header/footer
hp=sec.header.paragraphs[0]; hp.alignment=WD_ALIGN_PARAGRAPH.RIGHT; r=hp.add_run('AQUABOX | PLANEJAMENTO DE REQUISITOS'); set_font(r,8.5,True,'666666')
fp=sec.footer.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=fp.add_run('Documento de trabalho - Firmware Aquabox'); set_font(r,8,color='666666')

# Cover
p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(70); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('AQUABOX'); set_font(r,30,True,NAVY)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(20)
r=p.add_run('Planejamento de Requisitos do Firmware'); set_font(r,18,False,DARK)
add_callout(doc,'PROPÓSITO','Definir o escopo funcional, técnico e de segurança do firmware que automatiza o enchimento de caixas d\'água e a irrigação de jardim, usando ESP32-S3 como plataforma de controle.')
meta=[('Versão','0.1 - rascunho inicial'),('Data','02 de setembro de 2026'),('Público','Desenvolvimento de firmware, hardware, testes e produto'),('Status','Base para validação e detalhamento técnico')]
add_table(doc,['CAMPO','VALOR'],meta,[2700,6660])
doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

add_heading(doc,'1. Visão geral e objetivo',1)
doc.add_paragraph('O Aquabox é um controlador autônomo para gestão de água residencial ou de pequeno porte. O equipamento deve comandar até três válvulas solenoides e uma bomba, permitindo que cada válvula seja configurada individualmente para abastecimento de caixa d\'água ou irrigação. A operação normal é automática, com modos manuais físicos para manutenção e operação independente da bomba.')
add_heading(doc,'1.1 Objetivos do produto',2)
add_bullets(doc,[
    'Evitar falta de água e transbordamento ao controlar o enchimento de até três caixas d\'água.',
    'Executar rotinas de irrigação por horário e duração configuráveis.',
    'Oferecer interface local simples no display TFT, com quatro botões de navegação.',
    'Disponibilizar monitoramento e comandos remotos seguros por MQTT, sem substituir as proteções locais.',
    'Aumentar a segurança operacional por intertravamentos, alarmes e comportamento seguro diante de falhas.',
    'Manter data e hora para rotinas agendadas mesmo durante quedas de energia, por meio do RTC.'
])
add_heading(doc,'1.2 Fora de escopo desta versão',2)
add_bullets(doc,['Aplicativo móvel, conectividade Wi-Fi, telemetria em nuvem e atualização remota de firmware.','Medição volumétrica para cobrança, tratamento de água ou controle de fertilizantes.','Controle simultâneo de múltiplas bombas; esta especificação considera uma única saída de bomba.'])

add_heading(doc,'2. Interfaces e recursos de hardware',1)
add_table(doc,['ID','INTERFACE','FUNÇÃO NO FIRMWARE'],[
('HW-01','ESP32-S3','Microcontrolador dual-core; executa FreeRTOS, controle, interface de usuário, persistência, diagnóstico e conectividade Wi-Fi/MQTT.'),
('HW-02','TFT 1,8" ST7735','Exibe estado, alarmes, relógio, menus, configurações e comandos locais.'),
('HW-03','RTC DS3231M','Fornece data/hora para irrigação e registro de eventos; deve ser validado na inicialização.'),
('HW-04','Botões UP, DOWN, ENTER, BACK','Navegação, edição, confirmação e retorno nos menus.'),
('HW-05','Relé da bomba','Aciona a bomba d\'água conforme demanda autorizada ou chave física independente.'),
('HW-06','3 saídas de solenoide','Comandam abertura/fechamento de fluxo para caixa d\'água ou irrigação.'),
('HW-07','3 sensores de nível baixo','Indicam demanda de reabastecimento das caixas configuradas.'),
('HW-08','3 sensores de nível alto','Indicam nível máximo e encerram o enchimento das caixas configuradas.'),
('HW-09','Sensor de fluxo','Confirma fluxo durante operação da bomba e ajuda a detectar falha hidráulica.'),
('HW-10','Chave bomba independente','Solicita acionamento local da bomba, sujeito aos intertravamentos de segurança.'),
('HW-11','Chave Manual/Automático','Seleciona operação autônoma ou manutenção/limpeza com comandos manuais.')
],[1000,2100,6260])

add_heading(doc,'3. Modelo de configuração',1)
doc.add_paragraph('Cada canal de solenoide (S1, S2 e S3) deve possuir configuração independente. Um canal não utilizado deve poder ser desabilitado, permanecendo fechado e excluído das lógicas de demanda.')
add_table(doc,['PARÂMETRO POR CANAL','CAIXA D\'ÁGUA','IRRIGAÇÃO'],[
('Modo','Caixa d\'água','Irrigação'),('Habilitado','Sim/Não','Sim/Não'),('Vínculo de sensores','Sensor baixo e sensor alto do mesmo índice lógico','Não se aplica'),('Agenda','Não se aplica','Hora de início e duração; permitir pelo menos uma agenda por canal'),('Demanda de bomba','Solicita bomba enquanto a válvula estiver aberta','Solicita bomba durante o ciclo, salvo configuração futura diferente'),('Estado seguro','Fechar válvula em alarme crítico','Fechar válvula ao fim da duração ou em alarme crítico')
],[2500,3430,3430])
add_callout(doc,'REGRA DE VINCULAÇÃO','Na versão inicial, S1 deve usar os sensores de nível 1, S2 os sensores 2 e S3 os sensores 3 quando estiverem no modo Caixa d\'água. O mapeamento cruzado de sensores é uma evolução futura, para evitar complexidade de comissionamento.')

add_heading(doc,'4. Requisitos funcionais',1)
add_heading(doc,'4.1 Inicialização e estados gerais',2)
add_table(doc,['ID','REQUISITO','PRIORIDADE'],[
('RF-01','Ao energizar, o firmware deve manter bomba e todas as solenoides desligadas até concluir a leitura das entradas, do RTC e das configurações persistidas.','Alta'),
('RF-02','O firmware deve exibir uma tela de inicialização com versão do firmware e status resumido dos periféricos.','Média'),
('RF-03','O modo Manual/Automático efetivo deve ser lido da chave física e exibido permanentemente na tela principal.','Alta'),
('RF-04','Toda alteração de configuração confirmada pelo usuário deve ser gravada em memória não volátil.','Alta'),
('RF-05','Após reinicialização, o firmware deve restaurar configurações e retornar ao estado seguro, sem religar cargas automaticamente até reavaliar entradas.','Alta')
],[900,7050,1410])
add_heading(doc,'4.2 Controle de caixas d\'água',2)
add_table(doc,['ID','REQUISITO','PRIORIDADE'],[
('RF-10','Para canal em modo Caixa d\'água, a detecção de nível baixo deve criar demanda de enchimento quando o sistema estiver em Automático.','Alta'),
('RF-11','Com demanda válida, o firmware deve abrir a solenoide correspondente e solicitar a bomba.','Alta'),
('RF-12','A detecção de nível alto da caixa correspondente deve fechar a solenoide e remover sua demanda de bomba.','Alta'),
('RF-13','O estado de enchimento deve permanecer ativo entre nível baixo e nível alto, evitando oscilação por variações momentâneas do sensor.','Alta'),
('RF-14','Se uma caixa configurada apresentar baixo e alto ativos ao mesmo tempo por mais que o tempo de validação, o canal deve parar, gerar alarme de incoerência e requerer reconhecimento.','Alta'),
('RF-15','O firmware deve aplicar tempo máximo de enchimento configurável por canal; ao exceder, deve fechar a válvula, desligar a bomba se não houver outra demanda e gerar alarme.','Alta')
],[900,7050,1410])
add_heading(doc,'4.3 Controle de irrigação',2)
add_table(doc,['ID','REQUISITO','PRIORIDADE'],[
('RF-20','Para canal em modo Irrigação, o firmware deve abrir a solenoide no horário programado quando estiver em Automático.','Alta'),
('RF-21','O firmware deve manter a solenoide aberta pela duração configurada e fechá-la automaticamente ao término.','Alta'),
('RF-22','Cada agenda deve permitir habilitar/desabilitar o ciclo sem apagar horário ou duração.','Média'),
('RF-23','Quando dois ou mais canais de irrigação coincidirem, o firmware deve executar apenas um por vez, em ordem crescente de canal, para evitar exceder a capacidade hidráulica.','Alta'),
('RF-24','Uma irrigação pendente por conflito deve iniciar após a conclusão do ciclo anterior somente se ainda estiver dentro de uma janela configurável; fora da janela, deve ser registrada como perdida.','Média'),
('RF-25','O usuário deve poder interromper uma irrigação ativa pelo menu; a solenoide deve fechar imediatamente e o evento deve ser registrado.','Alta')
],[900,7050,1410])

add_heading(doc,'4.4 Bomba e sensor de fluxo',2)
add_table(doc,['ID','REQUISITO','PRIORIDADE'],[
('RF-30','A bomba deve ligar somente quando existir ao menos uma demanda autorizada de caixa ou irrigação, ou quando a chave independente solicitar operação manual.','Alta'),
('RF-31','Antes de acionar a bomba, o firmware deve abrir a solenoide solicitante e aguardar um atraso configurável de pré-abertura.','Alta'),
('RF-32','Após ligar a bomba, o firmware deve verificar a presença de pulsos do sensor de fluxo dentro de um tempo de partida configurável.','Alta'),
('RF-33','A ausência de fluxo durante o tempo de partida ou durante operação por tempo configurável deve desligar a bomba, fechar as válvulas ativas e gerar alarme de falta de fluxo.','Alta'),
('RF-34','O firmware deve calcular e exibir vazão aproximada e totalizador de pulsos/volume quando a constante de calibração for configurada.','Média'),
('RF-35','A chave de bomba independente deve ser indicada no display e não pode ignorar alarmes críticos de falta de fluxo, sobretempo ou falha de segurança.','Alta')
],[900,7050,1410])

add_heading(doc,'4.5 Operação manual e automática',2)
add_table(doc,['MODO','COMPORTAMENTO OBRIGATÓRIO'],[
('Automático','Executa enchimento por sensores e irrigação por agenda. Comandos manuais pela interface devem ser bloqueados, exceto reconhecimento de alarme e interrupção segura.'),
('Manual','Suspende o disparo automático de enchimento e irrigação. Permite abrir/fechar cada solenoide pelo menu para manutenção, sempre com confirmação visual. A bomba pode ser comandada pela chave independente, respeitando proteções.'),
('Transição para Manual','Ao mudar para Manual, as agendas não devem iniciar. Um ciclo automático em andamento deve ser encerrado de modo seguro: fechar válvulas e desligar a bomba após atraso de desligamento.'),
('Transição para Automático','O firmware deve reavaliar sensores e agenda; não deve iniciar saídas sem a sequência normal de autorização e pré-abertura.')
],[2100,7260])

add_heading(doc,'5. Interface de usuário',1)
add_table(doc,['TELA/MENU','CONTEÚDO E AÇÕES'],[
('Tela principal','Hora/data, modo Manual/Automático, estado da bomba, estado de S1/S2/S3, resumo de caixas e alarmes. ENTER abre o menu.'),
('Configuração de canal','Habilitar canal, selecionar Caixa/Irrigação, ajustar tempos máximos ou horário/duração conforme o modo.'),
('Configuração de sistema','Data/hora RTC, atrasos de acionamento, timeout de fluxo, constante de pulsos por litro, brilho do display e idioma futuro.'),
('Comandos manuais','Disponível somente em Manual: abrir/fechar solenoide e visualizar intertravamentos da bomba.'),
('Alarmes e diagnósticos','Lista de alarmes ativos, histórico resumido, estado bruto dos sensores, leitura de fluxo e informação de firmware.'),
('Navegação','UP/DOWN movem seleção ou ajustam valor; ENTER confirma/entra; BACK retorna/cancela. Pressão longa de BACK na tela principal deve reconhecer alarmes não críticos após confirmação.')
],[2400,6960])

add_heading(doc,'6. Regras de segurança e intertravamentos',1)
add_bullets(doc,[
    'Todas as saídas devem iniciar desligadas e permanecer desligadas na ausência de configuração válida ou em falha crítica.',
    'Nunca acionar bomba sem ao menos uma solenoide autorizada aberta, exceto em um modo de teste de comissionamento explicitamente protegido e ainda a definir.',
    'Uma falha de falta de fluxo, tempo máximo de enchimento, sensores incoerentes ou RTC inválido para irrigação deve impedir novo acionamento relacionado até reconhecimento e condição normalizada.',
    'O firmware deve aplicar debounce/filtragem configurável a botões, chaves e sensores de nível para evitar transições falsas.',
    'As solenoides devem fechar antes do desligamento da bomba, respeitando atraso configurável, exceto em situação de corte de segurança que exige desligamento imediato.',
    'Em reinício inesperado, watchdog ou falha de leitura de configuração, o sistema deve retornar ao estado seguro e registrar a causa, quando possível.'
])
add_table(doc,['ALARME','CONDIÇÃO','AÇÃO AUTOMÁTICA'],[
('AL-01 - Falta de fluxo','Bomba ligada sem pulsos válidos no tempo de partida/operação.','Desligar bomba, fechar válvulas ativas e bloquear reinício até reconhecimento.'),
('AL-02 - Tempo máximo','Canal de caixa não atingiu nível alto no limite configurado.','Fechar válvula do canal; desligar bomba se não houver outra demanda.'),
('AL-03 - Sensores incoerentes','Nível baixo e alto ativos simultaneamente após filtro.','Fechar válvula do canal, bloquear canal e sinalizar manutenção.'),
('AL-04 - RTC inválido','RTC ausente, leitura inválida ou horário não configurado.','Bloquear irrigação automática; manter enchimento disponível e avisar usuário.'),
('AL-05 - Configuração inválida','Parâmetros fora de faixa, corrupção ou combinação impossível.','Desabilitar canal afetado e solicitar reconfiguração.')
],[2100,3300,3960])

add_heading(doc,'7. Máquina de estados proposta',1)
add_table(doc,['ESTADO','ENTRADA/CONDIÇÃO','AÇÃO','SAÍDA'],[
('INICIALIZANDO','Energização ou reset','Carrega configuração, inicializa periféricos, lê entradas.','PRONTO ou FALHA SEGURA'),
('PRONTO','Automático, sem demanda','Monitora sensores, agenda, chaves e alarmes.','ENCHENDO, IRRIGANDO, MANUAL ou ALARME'),
('ENCHENDO','Demanda de caixa validada','Abre válvula, aguarda pré-abertura, liga bomba e monitora fluxo.','PRONTO ao nível alto; ALARME em falha'),
('IRRIGANDO','Agenda válida liberada','Abre válvula, liga bomba e conta duração.','PRONTO ao fim; ALARME em falha'),
('MANUAL','Chave em Manual','Bloqueia automações; aceita comandos locais permitidos.','PRONTO ao retornar Automático'),
('ALARME','Falha crítica ou bloqueante','Desliga saídas afetadas, informa causa e exige reconhecimento.','PRONTO/MANUAL quando condição normalizada')
],[1700,2800,2900,1960])

add_heading(doc,'8. Arquitetura de firmware com FreeRTOS',1)
doc.add_paragraph('O firmware deve utilizar FreeRTOS, disponível no ESP-IDF, para isolar atividades de tempo crítico das tarefas de conectividade e interface. O uso dos dois núcleos deve favorecer previsibilidade do controle hidráulico: a comunicação de rede não pode atrasar a leitura de sensores nem o desligamento de segurança.')
add_table(doc,['TAREFA','NÚCLEO PREFERENCIAL','RESPONSABILIDADE','CLASSE DE PRIORIDADE'],[
('Controle hidráulico','Core 1','Máquina de estados, leitura validada de nível, acionamento de válvulas/bomba, timeouts e intertravamentos.','Alta'),
('Aquisição de fluxo','Core 1 / ISR curta','Contagem de pulsos por interrupção; processamento e cálculo de vazão em tarefa.','Alta'),
('Wi-Fi e MQTT','Core 0','Conexão de rede, TLS, publicação de telemetria e recepção de comandos.','Média'),
('Interface local','Core 1','Leitura de botões/chaves, atualização do TFT e navegação de menus.','Média'),
('Persistência e log','Core 0','Gravação assíncrona de configuração e eventos, limitada para preservar a memória flash.','Baixa'),
('Supervisão','Core 1','Watchdog de tarefas, monitoramento de filas e sinais de falha entre tarefas.','Alta')
],[1900,1700,4100,1660])
add_heading(doc,'8.1 Comunicação e sincronização entre tarefas',2)
add_bullets(doc,[
    'A tarefa de Controle hidráulico deve ser a única autorizada a alterar as saídas de bomba e solenoides.',
    'Eventos de sensores, chaves, agenda e comandos MQTT devem chegar ao controle por filas FreeRTOS; callbacks e ISR não podem acionar cargas diretamente.',
    'Mensagens entre tarefas devem usar estruturas versionadas e filas com tamanho definido; eventos críticos precisam ter fila reservada ou mecanismo que impeça perda silenciosa.',
    'Mutexes devem proteger apenas recursos compartilhados de curta duração, como barramento I2C, display e armazenamento; a tarefa de controle não deve aguardar rede ou escrita em flash.',
    'Notificações de tarefa ou event groups devem sinalizar alarmes e mudanças de estado com baixa latência.'
])
add_heading(doc,'8.2 Requisitos de execução em tempo real',2)
add_table(doc,['ID','REQUISITO','PRIORIDADE'],[
('RNF-11','A tarefa de Controle hidráulico deve ter prioridade superior às tarefas MQTT, interface e persistência; ela não pode executar operações de rede, TLS ou gravação em flash.','Alta'),
('RNF-12','A leitura de pulsos do sensor de fluxo deve usar interrupção ou periférico apropriado e manter o tratamento de interrupção curto, sem alocação dinâmica ou chamadas bloqueantes.','Alta'),
('RNF-13','A recepção de comando MQTT deve apenas validar sintaticamente e enfileirar a solicitação; a autorização final e o acionamento devem ocorrer na tarefa de Controle hidráulico.','Alta'),
('RNF-14','O firmware deve monitorar uso de pilha, ocupação de filas e reinicializações de tarefas; condições fora do limite devem gerar diagnóstico e, se necessário, estado seguro.','Alta'),
('RNF-15','Temporizações de irrigação, pré-abertura e segurança devem usar relógio monotônico/timers do FreeRTOS, enquanto o RTC DS3231M deve ser usado para agendamento por data e hora.','Alta'),
('RNF-16','A divisão de tarefas e afinidade de núcleos deve ser documentada em código e testada sob carga de Wi-Fi/MQTT, garantindo que a lógica de segurança permaneça responsiva.','Alta')
],[900,7050,1410])
add_callout(doc,'PRINCÍPIO DE SEGURANÇA','FreeRTOS melhora a organização e o paralelismo, mas não muda a autoridade do controle local: comandos MQTT, botões e agenda devem convergir para a mesma máquina de estados e para os mesmos intertravamentos.')

add_heading(doc,'9. Monitoramento remoto e MQTT',1)
doc.add_paragraph('O Aquabox deve usar MQTT sobre Wi-Fi para publicar telemetria e receber comandos. A comunicação remota complementa a interface local; ela não pode contornar chaves físicas, intertravamentos nem alarmes críticos.')
add_heading(doc,'8.1 Conexão e tópicos',2)
add_table(doc,['ID','REQUISITO','PRIORIDADE'],[
('RF-40','O firmware deve permitir configurar SSID, senha Wi-Fi, endereço/porta do broker MQTT, credenciais, ID do dispositivo e prefixo de tópicos.','Alta'),
('RF-41','A conexão MQTT deve usar TLS quando o broker estiver configurado para conexão segura; certificados ou fingerprint devem ser persistidos de forma protegida.','Alta'),
('RF-42','O dispositivo deve publicar disponibilidade via mensagem de nascimento e testamento (LWT), indicando online/offline.','Alta'),
('RF-43','O firmware deve reconectar automaticamente ao Wi-Fi e ao broker com retentativas progressivas, sem bloquear a lógica local de controle.','Alta'),
('RF-44','A ausência de conexão MQTT não deve interromper a automação local nem impedir comandos pelos botões e chaves.','Alta')
],[900,7050,1410])
add_heading(doc,'8.2 Telemetria publicada',2)
add_table(doc,['TÓPICO RELATIVO','RETAIN','CONTEÚDO MÍNIMO'],[
('aquabox/{id}/status','Sim','Disponibilidade, modo Manual/Automático, versão, RSSI, hora, estado da bomba e alarmes ativos.'),
('aquabox/{id}/telemetry','Não','Leituras de nível baixo/alto, estado das solenoides, fluxo, vazão, totalizador, ciclo ativo e tempo restante.'),
('aquabox/{id}/event','Não','Inicialização, mudanças de modo, início/fim de ciclo, comandos recebidos, falhas e reconhecimentos.'),
('aquabox/{id}/config','Não','Resposta de leitura de configuração não sigilosa e confirmação de atualização aceita/rejeitada.')
],[3150,900,5310])
add_heading(doc,'8.3 Comandos remotos e segurança',2)
add_table(doc,['ID','REQUISITO','PRIORIDADE'],[
('RF-45','O firmware deve assinar apenas o tópico aquabox/{id}/command e validar o formato, versão, ID de correlação, timestamp e prazo de expiração de cada comando.','Alta'),
('RF-46','Cada comando deve retornar confirmação no tópico aquabox/{id}/response com ID de correlação, resultado, motivo de recusa e estado final observado.','Alta'),
('RF-47','Comandos remotos permitidos: leitura de estado/configuração, atualização de agenda/parâmetros não críticos, iniciar/parar irrigação e reconhecer alarme não crítico.','Alta'),
('RF-48','Comandos remotos que acionem bomba ou solenoide devem ser aceitos somente com a chave em Automático, sem alarme bloqueante, com intertravamentos satisfeitos e duração limitada configurável.','Alta'),
('RF-49','Comandos remotos nunca devem alterar o estado da chave Manual/Automático, desabilitar proteção de fluxo, ignorar nível alto ou desbloquear alarmes críticos.','Alta'),
('RF-50','O firmware deve rejeitar comandos duplicados, expirados, malformados ou de origem não autenticada e registrar a rejeição no log de eventos.','Alta'),
('RF-51','Configurações sigilosas, como senha Wi-Fi, senha MQTT e material de certificado, nunca devem ser publicadas em tópicos MQTT nem exibidas integralmente no display.','Alta')
],[900,7050,1410])
add_callout(doc,'EXEMPLO DE COMANDO','Publicar em aquabox/{id}/command: {"id":"cmd-123","action":"start_irrigation","channel":2,"duration_s":600,"expires_at":"2026-09-02T15:30:00Z"}. A execução depende dos mesmos intertravamentos aplicados localmente.')

add_heading(doc,'10. Requisitos não funcionais',1)
add_table(doc,['ID','REQUISITO'],[
('RNF-01','O ciclo principal deve responder a uma mudança de sensor de nível validada em até 1 segundo, excluído o tempo de filtro configurado.'),
('RNF-02','A interface local deve permanecer navegável durante monitoramento de fluxo e temporizações, sem bloqueios perceptíveis.'),
('RNF-03','Parâmetros críticos devem ter validação de faixa, valores padrão seguros e confirmação antes de gravar.'),
('RNF-04','O firmware deve separar camada de hardware, lógica de controle, persistência, interface e diagnóstico para permitir testes unitários.'),
('RNF-05','Eventos importantes devem ser armazenados em log circular: inicialização, acionamentos, encerramentos, alarmes, mudanças de configuração e alternância de modo.'),
('RNF-06','O projeto deve prever watchdog, tratamento de erro de periféricos e mecanismo de recuperação sem acionar cargas indevidamente.'),
('RNF-07','A versão do firmware, o esquema de configuração e a causa do último reset devem estar disponíveis no menu de diagnóstico.'),
('RNF-08','A telemetria MQTT deve usar payload JSON versionado, com limite de tamanho e taxa de publicação configurável para evitar saturar rede ou broker.'),
('RNF-09','Credenciais e parâmetros de rede devem ser armazenados em área não volátil com proteção contra leitura casual e não podem constar em logs.'),
('RNF-10','Falhas de Wi-Fi, DNS ou broker devem ser registradas com limitação de frequência, evitando desgaste excessivo da memória de log.')
],[1200,8160])

add_heading(doc,'11. Critérios de aceitação e testes',1)
add_table(doc,['CASO','CENÁRIO','RESULTADO ESPERADO'],[
('CT-01','Caixa 1: sensor baixo ativa, alto inativo, modo Automático.','S1 abre, bomba liga após pré-abertura, fluxo é detectado e a tela indica enchimento.'),
('CT-02','Durante CT-01, sensor alto da caixa 1 ativa.','S1 fecha; bomba desliga após atraso se não houver outra demanda.'),
('CT-03','Irrigação S2 programada para horário atual, duração de 2 min.','S2 abre e fecha após 2 min; bomba opera com fluxo válido.'),
('CT-04','S1 e S2 de irrigação iniciam no mesmo minuto.','Sistema executa um ciclo por vez na ordem de prioridade definida e registra o segundo como pendente.'),
('CT-05','Bomba ligada e sensor de fluxo sem pulsos.','Alarme AL-01 é emitido; bomba e válvulas ativas são desligadas.'),
('CT-06','Chave muda de Automático para Manual durante irrigação.','Ciclo é encerrado de forma segura; nenhuma nova agenda inicia em Manual.'),
('CT-07','Sensor baixo e alto ativos simultaneamente em uma caixa configurada.','Canal é bloqueado, válvula fecha e alarme AL-03 é exibido.'),
('CT-08','RTC desconectado ou data/hora inválida.','Irrigação automática fica bloqueada; enchimento por nível continua disponível; alarme AL-04 é exibido.'),
('CT-09','Reinício durante uma saída ativa.','Após reiniciar, todas as saídas permanecem desligadas até a reavaliação completa.'),
('CT-10','Usuário altera tempo máximo de enchimento fora de faixa.','Interface rejeita o valor, informa a faixa permitida e mantém o valor anterior.'),
('CT-11','Broker MQTT indisponível durante operação local.','Automação local continua; o firmware retenta conexão sem travar interface ou controle.'),
('CT-12','Comando MQTT válido para iniciar irrigação S2 em Automático, sem alarmes.','Sistema valida prazo e intertravamentos, inicia S2 respeitando pré-abertura e publica confirmação.'),
('CT-13','Comando MQTT tenta ligar bomba com chave em Manual ou sem válvula autorizada.','Comando é recusado, nenhuma saída é alterada e o motivo é publicado em response/event.'),
('CT-14','Comando MQTT expirado ou repetido.','Comando é recusado, identificado pelo ID de correlação e registrado no histórico.'),
('CT-15','Publicação MQTT contínua e reconexões de Wi-Fi durante enchimento.','A tarefa de Controle mantém leitura de nível, fluxo e timeout dentro dos limites definidos; não há atraso no desligamento por nível alto.'),
('CT-16','Fila de comandos MQTT cheia ou comando malformado sob carga.','O comando é descartado de forma controlada, com evento de diagnóstico; as saídas mantêm o último estado seguro.'),
('CT-17','Tarefa de persistência lenta durante operação crítica.','A gravação ocorre de forma assíncrona e não bloqueia a tarefa de Controle nem o processamento do sensor de fluxo.')
],[900,4000,4460])

add_heading(doc,'12. Decisões pendentes para a próxima revisão',1)
add_bullets(doc,[
    'Definir características elétricas e lógica ativa dos sensores de nível, chaves, relé e drivers de solenoide (ativo alto/baixo, pull-up/pull-down e isolamento).',
    'Definir limites configuráveis: duração máxima de irrigação, timeout de enchimento, janela de atraso, debounce e faixas da vazão esperada.',
    'Confirmar se a irrigação terá dias da semana, múltiplos horários por canal e regras para pular ciclos em caso de falta de água.',
    'Definir comportamento da chave de bomba independente: somente solicitação manual com válvula aberta ou operação de teste com regra adicional.',
    'Definir retenção de configuração e de histórico, além da necessidade de bateria no RTC e política para horário inválido.',
    'Definir broker MQTT, política de credenciais, uso obrigatório de TLS, certificado de autoridade e processo de provisionamento Wi-Fi.',
    'Definir permissões de comando remoto por perfil/usuário, retenção de telemetria e integração prevista com painel ou aplicativo.',
    'Definir orçamento de CPU, tamanho de pilha, prioridades e afinidade definitiva de cada tarefa FreeRTOS após prototipação com Wi-Fi e display reais.',
    'Validar requisitos de proteção elétrica, caixa, grau de proteção ambiental e normas aplicáveis ao conjunto.'
])
add_callout(doc,'PRÓXIMO PASSO RECOMENDADO','Transformar este planejamento em uma especificação de hardware/IO e um diagrama de estados detalhado, antes de iniciar a implementação do firmware.')

doc.core_properties.title='Planejamento de Requisitos - Firmware Aquabox'
doc.core_properties.subject='Especificação inicial de requisitos de firmware'
doc.core_properties.author='Aquabox'
doc.save(OUT)
print(OUT)
